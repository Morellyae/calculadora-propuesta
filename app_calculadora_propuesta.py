# app_calculadora_propuesta.py
# Instalación requerida: pip install streamlit docx fpdf

import streamlit as st
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import os

# -----------------------------
# Configuración del logo y colores
# -----------------------------
LOGO_PATH = "logo.png"
COLOR_PRINCIPAL = "#1F5D3B"  # Verde oscuro
COLOR_SECUNDARIO = "#E5F0EA" # Verde claro para fondos

# -----------------------------
# Recetas de ejemplo
# -----------------------------
recetas = {
    "Torta Chocolate": {
        "porciones_base": 20,
        "ingredientes": {
            "Harina": 400, "Azúcar": 350, "Huevos": 6,
            "Mantequilla": 200, "Cacao en polvo": 50,
            "Levadura": 15, "Leche": 250
        },
        "costo_ingredientes": {
            "Harina": 0.002, "Azúcar": 0.003, "Huevos": 150,
            "Mantequilla": 0.01, "Cacao en polvo": 0.015,
            "Levadura": 0.03, "Leche": 0.002
        },
        "procedimiento": [
            "Precalentar horno a 180°C.",
            "Batir mantequilla con azúcar.",
            "Agregar huevos uno a uno.",
            "Incorporar harina, cacao y levadura.",
            "Añadir leche poco a poco.",
            "Hornear 40 minutos."
        ],
        "cremas": {
            "Ganache Chocolate": {
                "ingredientes": {"Chocolate": 200, "Crema": 150},
                "costo_ingredientes": {"Chocolate": 0.02, "Crema": 0.01}
            }
        }
    }
}

# Sustituciones veganas
sustituciones = {
    "Mantequilla": {"Aceite vegetal": 0.75},
    "Leche": {"Bebida vegetal": 1},
    "Huevos": {"Huevo vegano": 1}
}

# -----------------------------
# Funciones
# -----------------------------
def aplicar_sustituciones(ingredientes, sustituciones):
    nuevos = {}
    for ing, cant in ingredientes.items():
        if ing in sustituciones:
            for sub, factor in sustituciones[ing].items():
                nuevos[sub] = round(cant * factor, 2)
        else:
            nuevos[ing] = cant
    return nuevos

def escalar_receta(ingredientes, porciones_base, porciones_nuevas):
    factor = porciones_nuevas / porciones_base
    return {ing: round(cant * factor, 2) for ing, cant in ingredientes.items()}

def calcular_costo(ingredientes, costos, porciones):
    total = 0
    for ing, cant in ingredientes.items():
        if ing in costos:
            total += cant * costos[ing]
    costo_porcion = round(total / porciones, 2)
    return round(total, 2), costo_porcion

def generar_lista_compras(ingredientes, cremas=None):
    lista = ingredientes.copy()
    if cremas:
        for crema in cremas.values():
            for ing, cant in crema["ingredientes"].items():
                if ing in lista:
                    lista[ing] += cant
                else:
                    lista[ing] = cant
    return lista

def generar_docx(nombre_receta, receta, procedimiento, porciones, costo_total, costo_porcion, cremas=None):
    doc = Document()
    if os.path.exists(LOGO_PATH):
        doc.add_picture(LOGO_PATH, width=Inches(2))
    doc.add_heading(nombre_receta, level=1)
    doc.add_paragraph(f"Porciones: {porciones}")
    doc.add_paragraph(f"Costo total: ₡{costo_total}")
    doc.add_paragraph(f"Costo por porción: ₡{costo_porcion}")

    doc.add_heading("Ingredientes", level=2)
    lista_compra = generar_lista_compras(receta, cremas)
    for ing, cant in lista_compra.items():
        doc.add_paragraph(f"{ing}: {cant} g")

    if cremas:
        doc.add_heading("Cremas / Ganaches", level=2)
        for nombre, datos in cremas.items():
            doc.add_paragraph(f"{nombre}:")
            for ing, cant in datos["ingredientes"].items():
                doc.add_paragraph(f" - {ing}: {cant} g")

    doc.add_heading("Procedimiento", level=2)
    for paso in procedimiento:
        doc.add_paragraph(f"- {paso}")

    filename = nombre_receta.replace(" ", "_") + ".docx"
    doc.save(filename)
    return filename

def generar_pdf(nombre_receta, receta, procedimiento, porciones, costo_total, costo_porcion, cremas=None):
    pdf = FPDF()
    pdf.add_page()
    if os.path.exists(LOGO_PATH):
        pdf.image(LOGO_PATH, x=10, y=8, w=40)
        pdf.ln(35)

    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, nombre_receta, ln=True)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 8, f"Porciones: {porciones}", ln=True)
    pdf.cell(0, 8, f"Costo total: ₡{costo_total}", ln=True)
    pdf.cell(0, 8, f"Costo por porción: ₡{costo_porcion}", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Ingredientes", ln=True)
    pdf.set_font("Arial", "", 12)
    lista_compra = generar_lista_compras(receta, cremas)
    for ing, cant in lista_compra.items():
        pdf.cell(0, 6, f"{ing}: {cant} g", ln=True)

    if cremas:
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, "Cremas / Ganaches", ln=True)
        pdf.set_font("Arial", "", 12)
        for nombre, datos in cremas.items():
            pdf.cell(0, 6, f"{nombre}:", ln=True)
            for ing, cant in datos["ingredientes"].items():
                pdf.cell(0, 6, f" - {ing}: {cant} g", ln=True)

    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 8, "Procedimiento", ln=True)
    pdf.set_font("Arial", "", 12)
    for paso in procedimiento:
        pdf.multi_cell(0, 6, f"- {paso}")

    filename = nombre_receta.replace(" ", "_") + ".pdf"
    pdf.output(filename)
    return filename

# -----------------------------
# Streamlit Interfaz Profesional
# -----------------------------
st.set_page_config(page_title="Propuesta Gastronómica", layout="wide")
st.markdown(
    f"""
    <style>
    .stApp {{
        background-color: {COLOR_SECUNDARIO};
    }}
    .css-1d391kg {{
        color: {COLOR_PRINCIPAL};
        font-weight: bold;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

st.image(LOGO_PATH, width=200)
st.title("📑 Propuesta Gastronómica Interactiva")

# Selección de receta
receta_seleccionada = st.selectbox("Selecciona la receta:", list(recetas.keys()))

# Número de porciones
porciones = st.number_input("Número de porciones:", min_value=1, value=20)

# Opción vegana
opcion_vegana = st.radio("Tipo de receta:", ["Normal", "Vegana"])

# Botón de calcular
if st.button("Generar propuesta"):
    receta_base = recetas[receta_seleccionada]["ingredientes"]
    costos_base = recetas[receta_seleccionada]["costo_ingredientes"]
    procedimiento = recetas[receta_seleccionada]["procedimiento"]
    cremas = recetas[receta_seleccionada]["cremas"]

    # Escalado
    receta_escalada = escalar_receta(receta_base, recetas[receta_seleccionada]["porciones_base"], porciones)

    # Sustituciones veganas
    if opcion_vegana == "Vegana":
        receta_final = aplicar_sustituciones(receta_escalada, sustituciones)
    else:
        receta_final = receta_escalada

    # Calcular costos
    costo_total, costo_porcion = calcular_costo(receta_final, costos_base, porciones)

    # Secciones colapsables
    with st.expander("Ingredientes Escalados"):
        for ing, cant in receta_final.items():
            st.write(f"{ing}: {cant} g")

    if cremas:
        with st.expander("Cremas / Ganaches"):
            for nombre, datos in cremas.items():
                st.write(f"**{nombre}**")
                for ing, cant in datos["ingredientes"].items():
                    st.write(f"- {ing}: {cant} g")

    with st.expander("Procedimiento"):
        for paso in procedimiento:
            st.write(f"- {paso}")

    st.markdown(f"**Costo total:** ₡{costo_total}")
    st.markdown(f"**Costo por porción:** ₡{costo_porcion}")

    # Generar documentos
    docx_file = generar_docx(receta_seleccionada, receta_final, procedimiento, porciones, costo_total, costo_porcion, cremas)
    pdf_file = generar_pdf(receta_seleccionada, receta_final, procedimiento, porciones, costo_total, costo_porcion, cremas)

    st.success("✅ Propuesta generada")
    st.download_button("📄 Descargar Word", data=open(docx_file, "rb").read(), file_name=docx_file)
    st.download_button("📕 Descargar PDF", data=open(pdf_file, "rb").read(), file_name=pdf_file)
