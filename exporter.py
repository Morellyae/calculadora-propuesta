# exporter.py
# Funciones para exportar recetas a PDF y Word con logo, pie de pÃ¡gina y costos.

from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from io import BytesIO
from recetas import RECETAS # Importado para obtener el factor de escala base
from costos import COSTOS_UNITARIOS # Importado para referencia (aunque la lÃ³gica de cÃ¡lculo estÃ¡ en app.py)

# ==========================
# Cargar logo en memoria
# ==========================
try:
    with open("logo.png", "rb") as f:
        LOGO_DATA = BytesIO(f.read())
except FileNotFoundError:
    LOGO_DATA = None
    print("Advertencia: 'logo.png' no encontrado. Se exportarÃ¡ sin logo.")


# ==========================
# Clase personalizada PDF con pie de pÃ¡gina
# ==========================
class PDF(FPDF):
    def footer(self):
        self.set_y(-15)  # posiciÃ³n desde el fondo
        self.set_font("Times", "I", 8)
        self.cell(
            0,
            10,
            "Calculadora de PastelerÃ­a Profesional - Chef More's", # CORRECCIÃ“N FINAL: Se cambiÃ³ el apÃ³strofe curvo 'â€™' por el apÃ³strofe estÃ¡ndar "'"
            0,
            0,
            "C",
        )


# ==========================
# Exportar a PDF
# La funciÃ³n ahora acepta los argumentos de costo
# ==========================
def export_to_pdf(nombre_receta, ingredientes, porciones, notas, costo_total_receta, costo_por_porcion):
    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()

    FONT_NAME = "Times"

    # === HEADER: Logo y TÃ­tulo ===
    if LOGO_DATA:
        LOGO_DATA.seek(0)
        pdf.image(LOGO_DATA, x=10, y=10, w=25)  # logo mÃ¡s pequeÃ±o
        pdf.set_xy(40, 15)  # mover tÃ­tulo mÃ¡s abajo a la derecha
        pdf.set_font(FONT_NAME, "B", 18)
        pdf.cell(0, 10, nombre_receta, 0, 1, "L")
        pdf.ln(15)  # espacio extra para que no choque con ingredientes
    else:
        pdf.set_font(FONT_NAME, "B", 16)
        pdf.cell(0, 12, nombre_receta, 0, 1, "C")
        pdf.ln(10)

    pdf.set_font(FONT_NAME, "", 12)
    pdf.cell(0, 8, f"Porciones Calculadas: {porciones}", 0, 1)
    pdf.ln(5)

    # === Tabla de ingredientes ===
    pdf.set_font(FONT_NAME, "B", 12)
    pdf.set_fill_color(200, 220, 255)
    pdf.cell(80, 10, "Ingrediente", 1, 0, "C", 1)
    pdf.cell(40, 10, "Cantidad", 1, 0, "C", 1)
    pdf.cell(40, 10, "Unidad", 1, 1, "C", 1)

    pdf.set_font(FONT_NAME, "", 11)
    
    # Determinamos el factor de escala para mostrar las cantidades correctas en el PDF
    # NOTA: En la lÃ³gica de Streamlit (app.py) ya se hace este cÃ¡lculo
    # AquÃ­ solo mostramos el dato que viene en 'ingredientes' (que ya estÃ¡ escalado)
    for ing in ingredientes:
        cantidad_str = str(ing.get("cantidad_escalada", ing["cantidad"])) # Usamos el valor escalado de app.py
        
        pdf.cell(80, 8, ing["nombre"], 1, 0)
        pdf.cell(40, 8, cantidad_str, 1, 0, "R")
        pdf.cell(40, 8, ing["unidad"], 1, 1)


    # === SECCIÃ“N DE COSTOS ===
    if costo_total_receta > 0:
        pdf.ln(10)
        pdf.set_font(FONT_NAME, "B", 12)
        # El texto es simple ASCII para evitar problemas de codificaciÃ³n.
        pdf.cell(0, 8, "Costo de la Receta", 0, 1) 
        
        pdf.set_font(FONT_NAME, "", 11)
        pdf.cell(0, 7, f"Costo Total ({porciones} porciones): ${costo_total_receta:.2f}", 0, 1)
        pdf.cell(0, 7, f"Costo por PorciÃ³n: ${costo_por_porcion:.2f}", 0, 1)


    # === Notas ===
    if notas:
        pdf.ln(10)
        pdf.set_font(FONT_NAME, "B", 12)
        pdf.cell(0, 8, "Notas:", 0, 1)
        pdf.set_font(FONT_NAME, "", 10)
        pdf.multi_cell(0, 5, notas)

    # Exportar a memoria
    # Usamos BytesIO para el manejo robusto de bytes en Streamlit
    buffer = BytesIO()
    try:
        # Intentamos obtener el resultado directamente como bytes
        pdf_bytes = pdf.output(dest="S").encode("latin-1", "replace")
    except UnicodeEncodeError:
        # Fallback si hay problemas de codificaciÃ³n de caracteres especiales
        pdf_bytes = pdf.output(dest="S").encode("latin-1", "ignore")
        
    buffer.write(pdf_bytes)
    return buffer.getvalue()


# ==========================
# Exportar a DOCX
# La funciÃ³n ahora acepta los argumentos de costo
# ==========================
def export_to_docx(nombre_receta, ingredientes, porciones, notas, costo_total_receta, costo_por_porcion):
    doc = Document()

    # === HEADER: Logo y TÃ­tulo ===
    if LOGO_DATA:
        LOGO_DATA.seek(0)
        logo_copy = BytesIO(LOGO_DATA.read())
        doc.add_picture(logo_copy, width=Inches(0.75))
        doc.add_heading(nombre_receta, 0)
    else:
        doc.add_heading(nombre_receta, 0)

    doc.add_paragraph(f"Porciones calculadas: {porciones}")
    doc.add_paragraph()
    
    # --- SECCIÃ“N DE COSTOS ---
    if costo_total_receta > 0:
        # El DOCX puede manejar el emoji sin problemas
        doc.add_heading("ðŸ’° Costo de la Receta", level=1)
        doc.add_paragraph(f"Costo Total ({porciones} porciones): ${costo_total_receta:.2f}")
        doc.add_paragraph(f"Costo por PorciÃ³n: ${costo_por_porcion:.2f}")
        doc.add_paragraph()


    # Ingredientes
    doc.add_heading("Ingredientes", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Shading Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Ingrediente"
    hdr_cells[1].text = "Cantidad"
    hdr_cells[2].text = "Unidad"

    for ing in ingredientes:
        cantidad_str = str(ing.get("cantidad_escalada", ing["cantidad"])) # Usamos el valor escalado de app.py
        
        row_cells = table.add_row().cells
        row_cells[0].text = ing["nombre"]
        row_cells[1].text = cantidad_str
        row_cells[2].text = ing["unidad"]

    # Notas
    if notas:
        doc.add_heading("Notas", level=1)
        doc.add_paragraph(notas)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
